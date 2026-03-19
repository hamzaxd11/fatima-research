"""
Convert Markdown to a formatted DOCX file.

This converter is intentionally lightweight and dependency-minimal:
- Parses headings, paragraphs, bullet/numbered lists, tables, code blocks, and images
- Preserves basic inline formatting (**bold**, *italic*, `code`)
- Applies manuscript-friendly defaults (Times New Roman, 12pt, 1.5 spacing)

Usage:
    python scripts/markdown_to_docx.py RESEARCH_PAPER_RAW.md RESEARCH_PAPER.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def configure_document(doc: Document) -> None:
    """Apply document-wide formatting defaults."""
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.font.size = Pt(16)
    title.font.color.rgb = RGBColor(0, 0, 0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 4"].font.size = Pt(12)


def add_page_number(section) -> None:
    """Insert dynamic page number into section footer."""
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def parse_inline_tokens(text: str) -> List[Tuple[str, str]]:
    """Split text into plain/bold/italic/code segments."""
    parts: List[Tuple[str, str]] = []
    last = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > last:
            parts.append(("plain", text[last : match.start()]))
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            parts.append(("bold", token[2:-2]))
        elif token.startswith("*") and token.endswith("*"):
            parts.append(("italic", token[1:-1]))
        elif token.startswith("`") and token.endswith("`"):
            parts.append(("code", token[1:-1]))
        last = match.end()
    if last < len(text):
        parts.append(("plain", text[last:]))
    return parts


def add_inline_runs(paragraph, text: str) -> None:
    """Render inline markdown formatting inside a paragraph."""
    for kind, value in parse_inline_tokens(text):
        run = paragraph.add_run(value)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_separator_row(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def parse_table_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def read_table(lines: List[str], start: int) -> Tuple[List[List[str]], int]:
    table_lines: List[str] = []
    i = start
    while i < len(lines) and is_table_line(lines[i]):
        table_lines.append(lines[i])
        i += 1

    rows = [parse_table_row(x) for x in table_lines]
    if len(rows) >= 2 and is_separator_row(table_lines[1]):
        rows.pop(1)
    return rows, i


def add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(3)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if r_idx == 0:
                    run.bold = True
            # Right-align mostly numeric cells (except header row)
            if r_idx > 0 and re.fullmatch(r"[-+]?\d[\d,]*(\.\d+)?%?", text.replace(" ", "")):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()


def is_block_starter(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return True
    return bool(
        re.match(r"^(#{1,6})\s+", stripped)
        or re.match(r"^\d+\.\s+", stripped)
        or re.match(r"^-\s+", stripped)
        or re.match(r"^!\[.*\]\(.*\)$", stripped)
        or stripped in {"---", "***", "___"}
        or stripped.startswith("```")
        or is_table_line(stripped)
    )


def convert_markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    text = markdown_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])

    first_h1_used_as_title = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            code_lines: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            # Skip closing ``` if present
            if i < len(lines) and lines[i].strip().startswith("```"):
                i += 1

            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(cl)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            doc.add_paragraph()
            continue

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()

            if level == 1 and not first_h1_used_as_title:
                p = doc.add_paragraph(title_text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_h1_used_as_title = True
            else:
                heading_level = min(level, 4)
                doc.add_heading(title_text, level=heading_level)
            i += 1
            continue

        # Horizontal rule
        if stripped in {"---", "***", "___"}:
            doc.add_paragraph()
            i += 1
            continue

        # Image
        img_match = re.match(r"^!\[(.*)\]\((.*)\)$", stripped)
        if img_match:
            caption = img_match.group(1).strip()
            img_ref = img_match.group(2).strip()
            img_path = Path(img_ref)
            if not img_path.is_absolute():
                img_path = (markdown_path.parent / img_path).resolve()

            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.2))
            else:
                p = doc.add_paragraph(f"[Image not found: {img_ref}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

            if caption:
                cp = doc.add_paragraph(caption)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cp.runs:
                    run.italic = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)

            doc.add_paragraph()
            i += 1
            continue

        # Table
        if is_table_line(stripped):
            rows, i = read_table(lines, i)
            add_table(doc, rows)
            continue

        # Bullet list item
        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, bullet_match.group(1).strip())
            i += 1
            continue

        # Numbered list item
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, number_match.group(1).strip())
            i += 1
            continue

        # Plain paragraph (merge wrapped lines until next block)
        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip() or is_block_starter(nxt):
                break
            paragraph_lines.append(nxt.strip())
            i += 1

        paragraph_text = " ".join(paragraph_lines)
        p = doc.add_paragraph()
        if paragraph_text.startswith("**Supervisor:**") or paragraph_text.startswith("**Investigators:**"):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline_runs(p, paragraph_text)

    # Start references on a new page if that heading exists
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() in {"11. references", "references"}:
            if idx > 0:
                prev = doc.paragraphs[idx - 1]
                if prev.runs:
                    prev.runs[-1].add_break(WD_BREAK.PAGE)
                else:
                    prev.add_run().add_break(WD_BREAK.PAGE)
            break

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Markdown manuscript to formatted DOCX")
    parser.add_argument("input_markdown", help="Path to input markdown file")
    parser.add_argument(
        "output_docx",
        nargs="?",
        help="Path to output docx file (default: same name as input with .docx)",
    )
    args = parser.parse_args()

    input_path = Path(args.input_markdown).resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"Input markdown file not found: {input_path}")

    if args.output_docx:
        output_path = Path(args.output_docx).resolve()
    else:
        output_path = input_path.with_suffix(".docx")

    convert_markdown_to_docx(input_path, output_path)
    print(f"DOCX generated: {output_path}")


if __name__ == "__main__":
    main()
