"""Validate humanized manuscript variants against the corrected source paper."""

import hashlib
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
STEMS = (
    "RESEARCH_PAPER - PLAIN SCHOLARLY",
    "RESEARCH_PAPER - NATURAL ACADEMIC",
    "RESEARCH_PAPER - CONCISE JOURNAL",
)
CRITICAL_TEXT = (
    "H = 8.0427",
    "p = 0.0900",
    "H = 10.1562",
    "p = 0.0379",
    "0.0758",
    "p = 0.0705",
    "0.0676",
    "0.0853",
)
CRITICAL_PATTERNS = (
    r"(?:40|Forty) fully blank trailing (?:records|rows)",
    r"550 (?:missing cells|cells were missing)",
    r"six (?:out-of-label|categorical responses (?:were outside|fell outside))",
    r"Ethics approval and parental consent documentation could not be verified",
    r"Practice(?: scores)? differed significantly across maternal education groups",
    r"conservative exploratory Holm",
)
DISPLAY_LIMITS = {
    "RESEARCH_PAPER - PLAIN SCHOLARLY": (6, 4),
    "RESEARCH_PAPER - NATURAL ACADEMIC": (6, 4),
    "RESEARCH_PAPER - CONCISE JOURNAL": (3, 2),
}
FORBIDDEN_FRAMING = (
    "nominally significant",
    "did not survive Holm",
    "confirmatory result",
    "not familywise-error robust",
)
FORBIDDEN_INTERNAL_LANGUAGE = (
    "doc.md",
    "source protocol",
    "project workspace",
    "output/analysis_",
    "local analysis workflow",
    "reproducible pipeline",
    "timestamped outputs",
    "archived bundle",
    "appended consent text",
)


def extract_tables(text: str) -> list[str]:
    lines = text.splitlines()
    tables = []
    index = 0
    while index < len(lines):
        if not lines[index].startswith("|"):
            index += 1
            continue
        block = []
        while index < len(lines) and lines[index].startswith("|"):
            block.append(lines[index])
            index += 1
        tables.append("\n".join(block))
    return tables


def extract_dois(text: str) -> set[str]:
    matches = re.findall(r"(?:https://doi\.org/|doi:)(10\.\d{4,9}/[^\s]+)", text, re.IGNORECASE)
    return {match.rstrip(".,;)") for match in matches}


def section_text(text: str, start: str, end: str) -> str:
    return text.split(start, 1)[1].split(end, 1)[0]


def citation_order(text: str) -> list[int]:
    order = []
    for group in re.findall(r"\*\(([\d,–-]+)\)\*", text):
        for item in group.split(","):
            if "–" in item or "-" in item:
                start, end = re.split(r"[–-]", item)
                numbers = range(int(start), int(end) + 1)
            else:
                numbers = (int(item),)
            for number in numbers:
                if number not in order:
                    order.append(number)
    return order


def main() -> None:
    source = (ROOT / "RESEARCH_PAPER_RAW.md").read_text(encoding="utf-8")
    source_tables = extract_tables(source)
    source_dois = extract_dois(source)
    hashes = set()

    for stem in STEMS:
        markdown_path = ROOT / f"{stem}.md"
        docx_path = ROOT / f"{stem}.docx"
        text = markdown_path.read_text(encoding="utf-8")
        document = Document(docx_path)
        docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        flat_text = " ".join(text.split())
        flat_docx_text = " ".join(docx_text.split())
        expected_tables, expected_figures = DISPLAY_LIMITS[stem]
        word_details = ""

        references = re.findall(r"^(?:[1-9]|1[0-9]|2[01])\. ", text, re.MULTILINE)
        figures = re.findall(r"^!\[Figure \d+\.", text, re.MULTILINE)
        missing = [value for value in CRITICAL_TEXT if value.casefold() not in text.casefold()]
        missing_docx = [value for value in CRITICAL_TEXT if value.casefold() not in docx_text.casefold()]

        assert not missing, f"{stem}: missing critical Markdown text: {missing}"
        assert not missing_docx, f"{stem}: missing critical DOCX text: {missing_docx}"
        for pattern in CRITICAL_PATTERNS:
            assert re.search(pattern, flat_text, re.IGNORECASE), f"{stem}: missing Markdown pattern: {pattern}"
            assert re.search(pattern, flat_docx_text, re.IGNORECASE), f"{stem}: missing DOCX pattern: {pattern}"
        assert len(references) == 21, f"{stem}: expected 21 references"
        assert len(figures) == expected_figures, f"{stem}: expected {expected_figures} figure links"
        if stem != "RESEARCH_PAPER - CONCISE JOURNAL":
            assert extract_tables(text) == source_tables, f"{stem}: source tables changed"
        assert source_dois.issubset(extract_dois(text))
        assert "Excluded (missing maternal education)" not in text
        for phrase in FORBIDDEN_FRAMING:
            assert phrase not in text, f"{stem}: deprecated framing in Markdown: {phrase}"
            assert phrase not in docx_text, f"{stem}: deprecated framing in DOCX: {phrase}"
        for phrase in FORBIDDEN_INTERNAL_LANGUAGE:
            assert phrase not in text, f"{stem}: internal language in Markdown: {phrase}"
            assert phrase not in docx_text, f"{stem}: internal language in DOCX: {phrase}"
        assert "| Study | Main comparable finding |" not in text
        assert "### 6.2 Comparison with published studies" not in text
        assert "TODO" not in text and "[PLACEHOLDER]" not in text
        assert len(source_tables) == 6, "source manuscript should contain six result tables"
        assert len(extract_tables(text)) == expected_tables, f"{stem}: expected {expected_tables} Markdown tables"
        assert len(document.tables) == expected_tables, f"{stem}: expected {expected_tables} DOCX tables"
        assert len(document.inline_shapes) == expected_figures, f"{stem}: expected {expected_figures} DOCX figures"

        if stem == "RESEARCH_PAPER - CONCISE JOURNAL":
            title = text.splitlines()[0].removeprefix("# ")
            abstract = section_text(text, "## Abstract", "## Background")
            main_text = section_text(text, "## Background", "## Acknowledgements")
            word_details = f", {len(abstract.split())}-word abstract, {len(main_text.split())}-word main text"
            required_headings = ("## Background", "## Objectives", "## Methods", "## Results", "## Discussion", "## Conclusion")
            required_abstract_labels = ("**Background:**", "**Aim:**", "**Methods:**", "**Results:**", "**Conclusion:**")
            assert len(title.split()) <= 15, f"{stem}: title exceeds 15 words"
            assert all(heading in text for heading in required_headings), f"{stem}: missing EMHJ section"
            assert all(label in abstract for label in required_abstract_labels), f"{stem}: abstract is not structured"
            assert len(abstract.split()) <= 250, f"{stem}: abstract exceeds 250 words"
            assert len(main_text.split()) <= 3000, f"{stem}: main text exceeds 3000 words"
            assert expected_tables + expected_figures <= 5, f"{stem}: more than five display items"
            assert not re.search(r"\[\d+(?:[-,]\d+)*\]", text), f"{stem}: citations are not EMHJ parenthetical Vancouver style"
            assert re.search(r"\*\(1–3\)\*", text), f"{stem}: italic parenthetical citation not found"
            assert citation_order(text) == list(range(1, 22)), f"{stem}: references are not cited in first-appearance order"
            assert document.styles["Normal"].font.name == "Times New Roman"
            assert document.styles["Normal"].font.size.pt == 12
            assert document.styles["Normal"].paragraph_format.line_spacing == 2.0
            for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
                style = document.styles[style_name]
                assert style.font.name == "Times New Roman"
                assert style.font.size.pt == 12
                assert style.paragraph_format.line_spacing == 2.0
            assert all(
                run.font.size is not None and run.font.size.pt == 12
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
                for run in paragraph.runs
                if run.text
            ), f"{stem}: table text is not uniformly 12 point"
            citation_runs = [
                run
                for paragraph in document.paragraphs
                for run in paragraph.runs
                if re.fullmatch(r"\([\d,–-]+\)", run.text)
            ]
            assert citation_runs and all(run.italic for run in citation_runs), f"{stem}: DOCX citations are not italic"
            for section in document.sections:
                columns = section._sectPr.xpath("./w:cols")
                assert not columns or columns[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num") in (None, "1")

        digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
        hashes.add(digest)
        print(
            f"{stem}: {len(text.split())} words, 21 references, "
            f"{expected_tables} tables, {expected_figures} figures{word_details}, sha256={digest[:12]}"
        )

    assert len(hashes) == len(STEMS), "Variants are not textually distinct"
    print("ALL_VARIANTS_VERIFIED")


if __name__ == "__main__":
    main()
